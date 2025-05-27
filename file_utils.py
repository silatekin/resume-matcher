import logging
from docx import Document 
import pdfplumber
import fitz

def read_text_file(file_path):
    try:
        with open(file_path,'r',encoding='utf-8') as file:
            text = file.read()
        logging.info(f"Succesfully read file: {file_path}")
        return text
    except FileNotFoundError:
        logging.error(f"Error: File not found at {file_path}")
        return None
    except Exception as e:
        logging.error(f"Error reading file {file_path}:{e}")
        return None

def read_docx_file(file_path):
    
    full_text = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            full_text.append(para.text)
        logging.info(f"Succesfully read docx file from: {file_path}")
        return '\n'.join(full_text)  
    except FileNotFoundError:
        logging.error(f"Error: DOCX File not found at {file_path}")
        return None
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return None

def read_pdf_file(file_path):
    full_text = []
    try: 
        with fitz.open(file_path) as pdf:
            for page in pdf.pages:
                extracted_page_text = page.extract_text()
                if extracted_page_text:
                    full_text.append(extracted_page_text.strip()) 
        
        raw_text = '\n'.join(filter(None,full_text))

        if not raw_text.strip():
            logging.warning(f"No text extracted or PDF is empty: {file_path}")
            return None
        logging.info(f"Successfully read PDF file from: {file_path}")
        return raw_text
    except FileNotFoundError:
         logging.error(f"Error: PDF File not found at {file_path}")
         return None
    except Exception as e:
        logging.error(f"Error reading PDF file {file_path}: {e}")
        return None
    

def get_text_from_txt_object(uploaded_file_object):
    try:
        raw_text = uploaded_file_object.read().decode('utf-8')
        logging.info("Successfully read TXT file content.")
        return raw_text
    except Exception as e:
        logging.error(f"Error reading TXT file object: {e}")
        return None

def get_text_from_docx_object(uploaded_file_object):
    try:
        doc = Document(uploaded_file_object) 
        full_text = [para.text for para in doc.paragraphs]
        raw_text = '\n'.join(full_text)
        logging.info("Successfully read DOCX file content.")
        return raw_text
    except Exception as e:
        logging.error(f"Error reading DOCX file object: {e}")
        return None

def get_text_from_pdf_object_fitz(uploaded_file_object):
    full_text = []
    try:
        if hasattr(uploaded_file_object, 'seek') and callable(uploaded_file_object.seek):
            uploaded_file_object.seek(0)
        
        pdf_bytes = uploaded_file_object.read() 


        if hasattr(uploaded_file_object, 'seek') and callable(uploaded_file_object.seek):
             uploaded_file_object.seek(0) 

        with fitz.open(stream=pdf_bytes, filetype="pdf") as doc:
            for page_num in range(len(doc)):
                page = doc.load_page(page_num)
                extracted_page_text = page.get_text("text")
                if extracted_page_text: 
                    full_text.append(extracted_page_text.strip())
        
        raw_text = '\n'.join(filter(None, full_text))

        if not raw_text.strip():
            file_name_attr = getattr(uploaded_file_object, 'name', 'Uploaded PDF Object')
            logging.warning(f"PyMuPDF: No text extracted from PDF: {file_name_attr}")
            return None 
        logging.info(f"PyMuPDF: Successfully read PDF file content from object: {getattr(uploaded_file_object, 'name', '')}")
        return raw_text
    except Exception as e:
        file_name_attr = getattr(uploaded_file_object, 'name', 'Uploaded PDF Object')
        logging.error(f"PyMuPDF Error reading PDF file object {file_name_attr}: {e}")
        return None

def get_text_from_pdf_object(uploaded_file_object):
    try:
        full_text = []

        if hasattr(uploaded_file_object, 'seek') and callable(uploaded_file_object.seek):
            uploaded_file_object.seek(0)

        with pdfplumber.open(uploaded_file_object) as pdf:
            for page in pdf.pages:
                extracted_page_text = page.extract_text()
                if extracted_page_text: 
                    full_text.append(extracted_page_text)
        raw_text = '\n'.join(full_text)
        if not raw_text.strip():
            file_name_attr = getattr(uploaded_file_object, 'name', 'Uploaded PDF Object')
            logging.warning(f"No text extracted from PDF: {file_name_attr}")
            return None 
        logging.info("Successfully read PDF file content.")
        return raw_text
    except Exception as e:
        file_name_attr = getattr(uploaded_file_object, 'name', 'Uploaded PDF Object')
        logging.error(f"Error reading PDF file object {file_name_attr}: {e}")
        return None
