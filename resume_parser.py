import re
import spacy
import json
from collections import defaultdict
import datetime
from spacy.matcher import PhraseMatcher
from dateutil.parser import parse as parse_datetime
from dateutil.relativedelta import relativedelta
import logging
import os
from docx import Document
import pdfplumber

logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')

NLP_MODEL_GLOBAL = None 
try:
    NLP_MODEL_GLOBAL = spacy.load("en_core_web_md") 
    logging.info("Global NLP Model ('en_core_web_md') loaded successfully.")
except OSError:
    logging.error("Spacy model 'en_core_web_md' not found. Please run 'python -m spacy download en_core_web_md'")
    logging.info("Attempting to load 'en_core_web_sm' as a fallback...")
    try:
        NLP_MODEL_GLOBAL = spacy.load("en_core_web_sm")
        logging.info("Global NLP Model ('en_core_web_sm') loaded successfully as fallback.")
    except OSError:
        logging.error("Spacy model 'en_core_web_sm' also not found. NLP features will be unavailable if not passed explicitly.")


def read_text_file(file_path):
    try:
        with open(file_path,'r',encoding='utf-8') as file:
            text = file.read()
        logging.info(f"Succesfully read file: {file_path}")
        return text
    except FileNotFoundError:
        logging.error(f"Error: File not found at {file_path}")
        return None
    except Exception as e:
        logging.error(f"Error reading file {file_path}:{e}")
        return None

def read_docx_file(file_path):
    
    full_text = []
    try:
        doc = Document(file_path)
        for para in doc.paragraphs:
            full_text.append(para.text)
        logging.info(f"Succesfully read docx file from: {file_path}")
        return '\n'.join(full_text)  
    except FileNotFoundError:
        logging.error(f"Error: DOCX File not found at {file_path}")
        return None
    except Exception as e:
        logging.error(f"Error reading DOCX file {file_path}: {e}")
        return None

def read_pdf_file(file_path):
    full_text = []
    try: 
        with pdfplumber.open(file_path) as pdf:
            for page in pdf.pages:
                extracted_page_text = page.extract_text()
                if extracted_page_text:
                    full_text.append(extracted_page_text.strip()) 
        
        raw_text = '\n'.join(filter(None,full_text))

        if not raw_text.strip():
            logging.warning(f"No text extracted or PDF is empty: {file_path}")
            return None
        logging.info(f"Successfully read PDF file from: {file_path}")
        return raw_text
    except FileNotFoundError:
         logging.error(f"Error: PDF File not found at {file_path}")
         return None
    except Exception as e: # Catch other pdfplumber or general exceptions
        logging.error(f"Error reading PDF file {file_path}: {e}")
        return None


def clean_text(text):
    if not isinstance(text, str):
        logging.warning("clean_text received non-string input, returning empty string.")
        return ""
    text = text.strip()
    text = text.replace('\xa0', ' ') 
    text = '\n'.join([line.strip() for line in text.splitlines() if line.strip()])
    return text

def load_skills(skill_file=None):

    if skill_file is None:
        script_dir = os.path.dirname(os.path.abspath(__file__)) 
        skill_file = os.path.join(script_dir, 'tests', 'data', 'skills.json') 

    try:
        with open(skill_file,"r",encoding="utf-8") as f:
            skills_data = json.load(f)

            valid_skills = [str(skill).lower()
                            for skill in skills_data 
                            if isinstance(skill,str)]
            logging.info(f"Loaded {len(valid_skills)} skills from {skill_file}")
            return valid_skills
    except FileNotFoundError:
        logging.error(f"Skill file not found at {skill_file}") 
        return []
    except json.JSONDecodeError:
        logging.error(f"Error decoding JSON from skill file: {skill_file}") 
        return []
    except Exception as e:
        logging.error(f"General error loading skills from {skill_file}: {e}") 
        return []

TECH_SKILLS_LIST_GLOBAL = load_skills()
TECH_SKILLS_SET_GLOBAL = set(TECH_SKILLS_LIST_GLOBAL)

# do NOT use \\s
SECTION_HEADERS_GLOBAL = {
    "summary": r"^\s*(summary|profile|objective|about\s*me)([:\s]|\s*$)", # Changed \\s to \s
    "skills": r"^\s*(skills|technical\s*skills|technical\s*proficiency|core\s*competencies|technologies)([:\s]|\s*$)", # Changed \\s to \s
    "experience": r"^\s*(experience|work\s*experience|employment\s*history|professional\s*experience)([:\s]|\s*$)", # Changed \\s to \s
    "education": r"^\s*(education|academic\s*background|academic\s*qualifications)([:\s]|\s*$)", # Changed \\s to \s
    "projects": r"^\s*(projects|personal\s*projects)([:\s]|\s*$)", # Changed \\s to \s
}

EDUCATION_LEVELS_GLOBAL = {
    "phd": 5, "doctorate": 5, "dphil": 5,
    "master": 4, "masters": 4,
    "msc": 4, "m.sc": 4, "m.s.": 4, "ms": 4,
    "mba": 4, "m.b.a": 4,
    "meng": 4, "m.eng": 4,
    "mtech": 4, "m.tech": 4,
    "ma": 4, "m.a.": 4,
    "bachelor": 3, "bachelors": 3,
    "bsc": 3, "b.sc": 3, "b.s.": 3, "bs": 3,
    "beng": 3, "b.eng": 3,
    "btech": 3, "b.tech": 3,
    "ba": 3, "b.a.": 3,
    "bca": 3,
    "associate": 2, "associates": 2, "aa": 2, "as": 2,
    "college": 1,
    "high school": 0, "ged": 0, "diploma": 0
}

def segment_resume(raw_text, section_headers_dict): # section_headers_dict is kit.SECTION_HEADERS
    if not raw_text:
        return {}
    
    sections = {"header": ""}
    current_section_key = "header"
    lines = raw_text.splitlines()
    compiled_patterns = {}

    logging.debug("--- Compiling Section Header Regex Patterns ---")
    # This loop uses section_headers_dict which should be the complete one from main via ParsingKit
    for key,pattern_str in section_headers_dict.items(): # Renamed pattern to pattern_str for clarity
        try:
            compiled_patterns[key] = re.compile(pattern_str,re.IGNORECASE) # Use pattern_str here
            logging.debug(f"Compiled pattern for '{key}': {pattern_str}")
        except re.error as e:
            logging.error(f"Regex error in pattern for '{key}': {pattern_str} - {e}")
            continue # Skip this pattern if invalid
    
    current_section_content = []

    for i, line in enumerate(lines): 
        line_text_for_match = line.strip() 
        
        if not line_text_for_match: 
            if line.strip(): # If original line had content (e.g. just spaces, but not empty)
                 current_section_content.append(line) # Add original line to preserve formatting
            continue
            
        matched_key = None
        logging.debug(f"--- SEGMENTATION ATTEMPT ON Line {i}: Raw='{line}', Stripped='{line_text_for_match}' ---") # Shows the line being processed
        if not compiled_patterns:
            logging.debug("Segment_Debug: compiled_patterns dictionary is EMPTY!")
        for key, pattern_obj in compiled_patterns.items():
            logging.debug(f"Segment_Debug: Trying Key='{key}', Pattern='{pattern_obj.pattern}', Against='{line_text_for_match}'")
            match_result = pattern_obj.match(line_text_for_match)
            if match_result: 
                logging.info(f"Segment_Debug: SUCCESS - Matched Key='{key}' for line '{line_text_for_match}'") 
                matched_key = key
                break
            else:
                logging.debug(f"Segment_Debug: FAILED to match Key='{key}' for line '{line_text_for_match}'")
        
        if matched_key: 
            if current_section_content: 
                sections[current_section_key] = "\n".join(current_section_content).strip()
            current_section_key = matched_key
            # Start new section with the original line to preserve any leading/trailing spaces if it's part of content
            current_section_content = [line] 
        else: 
            if line.strip(): # Add line to content if it's not just an empty stripped line
                current_section_content.append(line)

    if current_section_content:
        sections[current_section_key] = "\n".join(current_section_content).strip()

    sections = {k: v for k, v in sections.items() if v.strip()} # Ensure sections are not just whitespace
    logging.info(f"Segmented resume into sections: {list(sections.keys())}")
    return sections

class ParsingKit:
    def __init__(self,nlp_model,tech_skills_list_ref, tech_skills_set_ref, section_headers_ref, education_levels_ref):
        self.nlp = nlp_model
        self.tech_skills = tech_skills_list_ref
        self.tech_skills_set = tech_skills_set_ref
        self.SECTION_HEADERS = section_headers_ref
        self.EDUCATION_LEVELS = education_levels_ref

        if self.nlp is None:
            logging.warning("ParsingKit initialized with nlp_model as None. NER features will be limited.")


    def parse_date(self, date_string):
        if not date_string or not isinstance(date_string,str):
            return None
        
        date_string = date_string.strip()

        if date_string.lower() in ["present","current","now","today","til date"]:
            return datetime.datetime.now().date() 
        try:
            dt_obj = parse_datetime(date_string, default=datetime.datetime(1, 1, 1), fuzzy=False) 
            return dt_obj.date()
        except (ValueError, OverflowError,TypeError):
            pass

        patterns = [
            r"(\d{4})",                 # Matches Year only (e.g., "2020")
            r"(\w+)\s+(\d{4})",         # Month Name + Year ("May 2021")
            r"(\d{1,2})/(\d{4})",       # MM/YYYY (e.g., "05/2022")
            r"(\d{1,2})-(\d{4})",       # MM-YYYY (e.g., "05-2022")
        ]

        #---Plan B(if try block fails)
        for pattern in patterns:
            match = re.search(pattern, date_string, re.IGNORECASE)
            if match:
                try:
                    if len(match.groups()) == 1: #only year matched(pattern 1)
                        return datetime.datetime(int(match.group(1)), 1, 1) #assume Jan 1st
                    elif len(match.groups()) == 2: #Month/Year MM/YYYY matched
                        month_str = match.group(1)
                        year_str = match.group(2)
                        try:
                            month = int(month_str)
                        except ValueError: #month is not a number ie May
                            # Try parsing the month name using dateutil again
                            try:
                                month_dt = parse_datetime(month_str, default=datetime.datetime(int(year_str), 1, 1))
                                month = month_dt.month #Gets the month number
                            except (ValueError, TypeError):
                                logging.warning(f"Could not parse month name: '{month_str}' in '{date_string}'")
                                continue # Invalid month name, try next regex pattern
                        if not (1 <= month <= 12): 
                            logging.warning(f"Invalid month '{month}' in '{date_string}'")
                            continue    
                        #Construct date assuming 1st day
                        return datetime.datetime(int(year_str), month, 1)
                except (ValueError, OverflowError):
                    logging.warning(f"Error parsing matched date parts from '{date_string}' with pattern '{pattern}'")
                    continue

        logging.warning(f"Could not parse date string: '{date_string}'")
        return None 



    def get_education_level(self,text):
        highest_level = -1
        if not text:
            return highest_level
        
        text_lower = text.lower()
        text_lower = re.sub(r'\b([a-z])\.', r'\1', text_lower)

        for keyword,level in self.EDUCATION_LEVELS.items():
            pattern = r'\b' + re.escape(keyword) + r'\b'
            if re.search(pattern, text_lower):
                highest_level = max(highest_level, level)

        logging.info(f"Determined highest education level found in text: {highest_level}")
        return highest_level


def parse_resume_sections(sections, kit: ParsingKit):

    if kit.nlp is None:
        logging.error("Spacy NLP model not loaded. Cannot perform detailed parsing.")
        return {
            "skills": ["NLP model not loaded"],
            "education_details": [], "contact_info": {}, "experience": [],
            "total_years_experience": 0.0, "education_level": -1,
            "raw_entities": defaultdict(list), "companies": []
        }

    parsed_resume = {
                "summary_text": None,
                "skills":[],
                "education_details":[],
                "contact_info":{},
                "experience":[],
                "total_years_experience":0.0,
                "education_level": -1,
                "raw_entities": defaultdict(list), 
                "companies": [],
    }

    if "summary" in sections:
        summary_text_content = sections["summary"]
        parsed_resume["summary_text"] = summary_text_content.strip() # Store the raw summary text
        logging.info(f"Extracted summary section text (length: {len(parsed_resume['summary_text'])}).")
    else:
        logging.info("No 'summary' section found.")


    #-------Skill Extraction--------
    if "skills" in sections:
        skills_text = sections["skills"]
        skills_doc = kit.nlp(skills_text)
        found_skills = set() 

        matcher = PhraseMatcher(kit.nlp.vocab, attr="LOWER")
        patterns = [kit.nlp.make_doc(skill) for skill in kit.tech_skills]
        matcher.add("TECH_SKILLS",patterns)
        matches = matcher(skills_doc)

        matched_indices = set() 
        for match_id,start,end in matches:
            span = skills_doc[start:end] 
            found_skills.add(span.text.strip())
            for i in range(start, end):
                matched_indices.add(i)

        # Token matching for single-word skills but only if the token wasn't already part of a phrase match
        for token in skills_doc:
            if token.i not in matched_indices: 
                if token.lemma_.lower() in kit.tech_skills_set:
                     # avoiding adding very short words unless specific known skills
                     if len(token.lemma_) > 1 or token.lemma_.lower() in ['c', 'r', 'go']:
                        found_skills.add(token.lemma_.lower()) 


        parsed_resume["skills"] = sorted(list(found_skills))
        logging.info(f"Extracted {len(parsed_resume['skills'])} unique skills.")
    
    #-------Education Extraction--------
    # Initialize in the main parsed_resume dictionary
    if "education_details" not in parsed_resume or not isinstance(parsed_resume.get("education_details"), list):
        parsed_resume["education_details"] = []
    if "education_level" not in parsed_resume:
        parsed_resume["education_level"] = -1

    if "education" in sections:
        education_text = sections["education"]
        # Get overall highest education level from the entire section text first
        highest_edu_level_found = kit.get_education_level(education_text)
        parsed_resume["education_level"] = highest_edu_level_found
        
        # Split the education section into lines for individual processing
        potential_degree_lines = [line.strip() for line in education_text.splitlines() if line.strip()]
        
        # Remove the "Education" header itself if it's the first line of the content
        if potential_degree_lines and potential_degree_lines[0].lower() == "education":
            potential_degree_lines = potential_degree_lines[1:]

        for line_text in potential_degree_lines:
            if not line_text: # Skip any remaining empty lines
                continue

            doc_line = kit.nlp(line_text) # Process this individual line with SpaCy

            # Initialize for this specific line/entry
            degree_mention = None
            institution_mention = None
            date_mention = None 
            
            line_lower_processed = line_text.lower().replace('\xa0', ' ') # Replace non-breaking space
            line_lower_processed = re.sub(r'\b([a-z])\.', r'\1', line_lower_processed) # Remove periods from B.S. -> bs

            # 1. Find Degree Mention on this line
            # Regex for specific degree patterns (e.g., B.S. in ..., Master of Science in ...)
            # This pattern tries to capture both abbreviations and fuller names with fields of study.
            specific_degree_pattern = re.compile(
                r"""
                \b(
                    (?:Bachelor|Master|Associate|Doctor(?:ate)?)\s*
                    (?:of|in|of\sScience|of\sArts|of\sBusiness\sAdministration|of\sEngineering|of\sPhilosophy)?\s* # Common "of X"
                    (?:in\s+)?[\w\s\(\)\-\.,'&]+?  # Field of study (allows more characters)
                    |                                   # OR
                    (?:B\.?S\.?C?|M\.?S\.?C?|M\.?B\.?A\.?|M\.A\.|A\.?A\.?S?|Ph\.?D\.?|B\.?Eng\.?|M\.?Eng\.?) # Abbreviations
                    (?:\s*(?:in|of)\s+[\w\s\(\)\-\.,'&]+?)? # Optional field for abbreviations
                )\b
                """, 
                re.IGNORECASE | re.VERBOSE
            )
            
            match = specific_degree_pattern.search(line_text)
            if match:
                extracted_degree_text = match.group(1).strip()
                # Clean up common trailing issues if regex is too greedy before a pipe or comma
                extracted_degree_text = re.sub(r'\s*[|].*$', '', extracted_degree_text).strip() # Remove from pipe
                extracted_degree_text = extracted_degree_text.rstrip(',').strip()
                # Further cleanup for abbreviations: B.S. -> BS, M.B.A. -> MBA
                extracted_degree_text = re.sub(r'\.(?=[A-Z])', '', extracted_degree_text.upper()).replace('.', '')
                degree_mention = extracted_degree_text
            else:
                # Fallback: if no specific pattern, check for keywords from EDUCATION_LEVELS
                # This helps if the format is very simple, e.g., just "Bachelor Degree"
                sorted_edu_keywords = sorted(kit.EDUCATION_LEVELS.items(), key=lambda item: len(item[0]), reverse=True)
                for keyword, level in sorted_edu_keywords:
                    pattern = r'\b' + re.escape(keyword) + r'\b'
                    if re.search(pattern, line_lower_processed):
                        degree_mention = keyword # Store the keyword (e.g., "bachelor")
                        break
            
            # 2. Find Institution on this line
            # Debug: Uncomment to see what SpaCy recognizes as ORG on this specific line
            # print(f"DEBUG Education Line: '{line_text}'")
            # print(f"DEBUG SpaCy Entities: {[(ent.text, ent.label_) for ent in doc_line.ents]}")

            potential_institutions_on_line = []
            for ent in doc_line.ents:
                if ent.label_ == "ORG":
                    ent_text_stripped = ent.text.strip()
                    # Filter out ORGs that might be part of the degree itself or common noise/headers
                    if degree_mention and ent_text_stripped.lower() in degree_mention.lower():
                        continue
                    if ent_text_stripped.lower() in ["education", "university", "college", "institute", "school"]:
                         if len(ent_text_stripped.split()) == 1: continue # Avoid generic terms if they are the whole ORG
                    
                    potential_institutions_on_line.append(ent_text_stripped)
            
            if potential_institutions_on_line:
                best_institution = ""
                for inst_text in potential_institutions_on_line:
                    # Prefer institutions that explicitly mention "college", "university", etc.
                    if any(kw in inst_text.lower() for kw in ["college", "university", "institute", "school"]):
                        if len(inst_text) > len(best_institution):
                            best_institution = inst_text
                
                if best_institution:
                    institution_mention = best_institution
                else: # If no keyword match, take the longest ORG entity found on the line
                    institution_mention = max(potential_institutions_on_line, key=len, default="") # default to "" if list empty
                
                institution_mention = institution_mention.rstrip(',').strip()
            
            # Fallback for institution if NER completely missed it (e.g., "BIGTOWN COLLEGE" in your example)
            if not institution_mention:
                # Look for text segments that are likely institutions based on keywords and capitalization
                # often found after the degree and date, possibly before a location
                parts_of_line = re.split(r'\s*\|\s*', line_text) # Split by pipe
                for part in reversed(parts_of_line): # Check from right to left
                    part = part.strip().rstrip(',')
                    if not part: continue
                    # Check for keywords or if it's a sequence of capitalized words (and not the date itself)
                    if date_mention and part.lower() == date_mention.lower(): continue
                    if degree_mention and part.lower() in degree_mention.lower() : continue

                    is_likely_inst = False
                    if any(kw in part.lower() for kw in ["college", "university", "institute", "school"]):
                        is_likely_inst = True
                    else: # Check for capitalized pattern if no keyword
                        words_in_part = part.split()
                        if len(words_in_part) > 0 and len(words_in_part) <= 4: # Reasonable length for an institution name
                            # Check if it's not a GPE (like CHICAGO, ILLINOIS)
                            doc_part_check = kit.nlp(part)
                            if not any(e.label_ == "GPE" and e.text == part for e in doc_part_check.ents):
                                if all(word[0].isupper() for word in words_in_part if word.lower() not in ["of", "the", "in", "at", "and"]):
                                    is_likely_inst = True
                    
                    if is_likely_inst:
                        # Try to remove city/state if appended with comma, e.g., "BIGTOWN COLLEGE, CHICAGO"
                        inst_candidate = re.sub(r',\s*(?:[A-Z]{2,}|[A-Za-z]+(?:,\s*[A-Z]{2})?)$', '', part).strip()
                        institution_mention = inst_candidate
                        break


            # 3. Find Date on this line (more targeted search)
            date_patterns_for_line = [
                r"((?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4})\b", # Month YYYY
                r"\b(\d{4})\b",                                  # YYYY (will be parsed as Jan 1st, YYYY)
                r"(\d{1,2}/\d{4})\b"                             # MM/YYYY
            ]
            for date_pattern_str in date_patterns_for_line:
                date_search_match = re.search(date_pattern_str, line_text, re.IGNORECASE)
                if date_search_match:
                    found_date_str = date_search_match.group(1)
                    # Try to parse this specific date string using the robust kit.parse_date
                    parsed_dt_object = kit.parse_date(found_date_str) 
                    if parsed_dt_object: # If parse_date successfully returns a datetime object
                        date_mention = found_date_str # Store the original string that was successfully parsed
                        break 
            
            # Append if we found at least a degree or an institution for this line
            if degree_mention or institution_mention:
                parsed_resume["education_details"].append({
                    "degree_mention": degree_mention, 
                    "institution_mention": institution_mention,
                    "date_mention": date_mention, 
                    "text": line_text # Store the original line text for this entry
                })
        
        logging.info(f"Extracted {len(parsed_resume['education_details'])} education details. Highest overall level: {highest_edu_level_found}")

    else: # No "education" section found in the segmented resume
        logging.info("No 'education' section found. Education level remains default.")
        # parsed_resume["education_level"] = -1 # Already initialized or set by get_education_level


    #-------Experience Extraction--------
    total_experience_duration_days = 0
    extracted_companies = set() # To store all unique company names found
    experience_text = sections.get("experience", "") 

    # Ensure parsed_resume["experience"] list exists
    if "experience" not in parsed_resume or not isinstance(parsed_resume.get("experience"), list):
        parsed_resume["experience"] = []

    # COMMON_SECTION_HEADERS is used to avoid misinterpreting section titles as job titles
    COMMON_SECTION_HEADERS = ["experience", "education", "skills", "summary", "objective",
                              "projects", "awards", "references", "publications", "interests",
                              "activities", "volunteer"] # Add any other common headers

    if experience_text:
        logging.info("Parsing experience from 'experience' section.")

        current_role = {}
        potential_header_lines = [] # Stores text of lines that might be part of a job header before a date is found

        lines = experience_text.splitlines()
        if lines: # Remove the "Experience" header itself if it's the first line of the section content
            first_line_cleaned_for_header_check = lines[0].strip().lower()
            # Check against various ways the "experience" header might be written
            if any(keyword in first_line_cleaned_for_header_check for keyword in ["experience", "work experience", "employment history", "professional experience"]):
                # More specific check: if it's JUST the header
                if len(first_line_cleaned_for_header_check.split()) <= 3: # "Experience" or "Work Experience"
                    lines = lines[1:]

        for line_content in lines:
            sent_text = line_content.strip()
            if not sent_text: # Skip empty lines
                continue

            # Regex to find date ranges (e.g., "Month YYYY - Month YYYY" or "Month YYYY - Present")
            date_range_pattern = r"(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|\d{4}|\d{1,2}/\d{4})\s*(?:-|–|to|until)\s*(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|\d{4}|\d{1,2}/\d{4}|Present|Current|Now)\b"
            date_match = re.search(date_range_pattern, sent_text, re.IGNORECASE)
            
            line_is_date = bool(date_match)

            if line_is_date:
                # Date line found: this anchors a new role or is part of the current role's header.
                # Finalize the PREVIOUS role if it was being built and had a start date.
                if current_role.get("start_date"):
                    prev_start = current_role.get("start_date")
                    prev_end = current_role.get("end_date")
                    if prev_start and prev_end:
                        try:
                            duration = relativedelta(prev_end, prev_start)
                            total_experience_duration_days += (duration.years * 365.25) + (duration.months * 30.4) + duration.days
                        except TypeError:
                            logging.warning(f"Could not calculate duration for role: {current_role.get('job_title') or current_role.get('company')}")
                    
                    # Append if it has at least a title, company or a valid date
                    if current_role.get("job_title") or current_role.get("company") or (prev_start and prev_end) :
                        parsed_resume["experience"].append(current_role)
                        logging.debug(f"===> Appended role (due to new date line): {current_role.get('job_title') or current_role.get('company')}")
                
                current_role = {} # Reset for the new role defined by this date line
                new_job_title = None
                new_company_name = None

                # 1. Parse dates from the current line
                start_date_str = date_match.group(1)
                end_date_str = date_match.group(2)
                start_date_obj = kit.parse_date(start_date_str)
                end_date_obj = kit.parse_date(end_date_str)
                
                description_from_date_line = ""
                text_before_date_on_date_line = sent_text[:date_match.start()].strip()
                text_after_date_on_date_line = sent_text[date_match.end():].strip()

                # If there's text after the date on the same line, it might be a description point
                if text_after_date_on_date_line:
                    potential_description = re.sub(r'^[-)*\u2022•·\s]+', '', text_after_date_on_date_line).strip()
                    if potential_description:
                        description_from_date_line = potential_description
                
                # Attempt to extract Title/Company from the text BEFORE the date on the current line
                if text_before_date_on_date_line:
                    parts = [p.strip() for p in text_before_date_on_date_line.split('|') if p.strip()]
                    
                    # Try to assign title and company based on pipe splitting
                    if len(parts) >= 1 : # Must have at least one part
                        # Heuristic: if two parts, assume Title | Company. If one part, could be Title or Company.
                        # Use NER to help disambiguate.
                        
                        part1_text = parts[0] if len(parts) > 0 else ""
                        part2_text = parts[1] if len(parts) > 1 else ""

                        doc_part1 = kit.nlp(part1_text) if part1_text else None
                        doc_part2 = kit.nlp(part2_text) if part2_text else None

                        is_part1_org = any(ent.label_ == "ORG" for ent in (doc_part1.ents if doc_part1 else []))
                        is_part2_org = any(ent.label_ == "ORG" for ent in (doc_part2.ents if doc_part2 else []))

                        if len(parts) == 2:
                            if is_part2_org and not is_part1_org: # Title | Company (ORG)
                                new_job_title = part1_text
                                new_company_name = part2_text
                            elif is_part1_org and not is_part2_org: # Company (ORG) | Title
                                new_company_name = part1_text
                                new_job_title = part2_text
                            elif is_part1_org and is_part2_org: # Both ORG, assume Title | Company
                                new_job_title = part1_text
                                new_company_name = part2_text
                            else: # Neither clearly ORG by NER, assume Title | Company
                                new_job_title = part1_text
                                new_company_name = part2_text
                        elif len(parts) == 1:
                            # Only one part. Could be title or company. Use NER.
                            if is_part1_org:
                                new_company_name = part1_text
                                # Title might be in potential_header_lines
                            else: # Assume it's a title
                                new_job_title = part1_text
                                # Company might be in potential_header_lines or next line

                    # If pipe splitting didn't yield company, try NER on the whole pre-date text
                    if not new_company_name:
                        doc_before_date = kit.nlp(text_before_date_on_date_line)
                        for ent in doc_before_date.ents:
                            if ent.label_ == "ORG":
                                # Avoid very long ORG entities that are likely misclassifications
                                if len(ent.text.split()) < 7:
                                    new_company_name = ent.text.strip()
                                    break # Take the first plausible ORG
                    
                    # If pipe splitting didn't yield title, but we got a company from NER
                    if not new_job_title and new_company_name:
                        company_idx = text_before_date_on_date_line.lower().find(new_company_name.lower())
                        if company_idx > 0: # Company name is not at the start
                            title_candidate = text_before_date_on_date_line[:company_idx].strip()
                            title_candidate = re.sub(r'[,|@\s]*at\s*$', '', title_candidate, flags=re.IGNORECASE).strip().rstrip(',').strip()
                            if title_candidate and (title_candidate[0].isupper() or title_candidate[0].isnumeric()):
                                new_job_title = title_candidate
                        # If company_idx is 0, the whole string was the company, title might be in potential_header_lines
                    
                    # If still no title and no company from pipe/NER on this line, but there was text
                    elif not new_job_title and not new_company_name and text_before_date_on_date_line:
                        if (text_before_date_on_date_line[0].isupper() or text_before_date_on_date_line[0].isnumeric()) and \
                           text_before_date_on_date_line.lower() not in COMMON_SECTION_HEADERS and \
                           len(text_before_date_on_date_line.split()) < 7: # Avoid long sentences as titles
                            new_job_title = text_before_date_on_date_line

                # 2. If title or company still missing, use potential_header_lines (lines before this date line)
                if not new_job_title or not new_company_name:
                    for header_line_text in reversed(potential_header_lines): # Process most recent first
                        if new_job_title and new_company_name: break # Both found

                        doc_header = kit.nlp(header_line_text)
                        temp_header_company = None
                        
                        if not new_company_name:
                            for ent in doc_header.ents:
                                if ent.label_ == "ORG":
                                    if len(ent.text.split()) < 7: # Avoid overly long ORGs
                                        new_company_name = ent.text.strip()
                                        temp_header_company = new_company_name # Mark company found on this line
                                        break
                        
                        if not new_job_title:
                            title_candidate_from_header = header_line_text # Assume whole line initially
                            if temp_header_company: # If company was found on THIS header_line
                                company_idx_h = header_line_text.lower().find(temp_header_company.lower())
                                if company_idx_h > 0:
                                    title_candidate_from_header = header_line_text[:company_idx_h].strip()
                                elif company_idx_h == 0: # Whole line was company name
                                    title_candidate_from_header = "" 
                            
                            title_candidate_from_header = re.sub(r'[,|@\s]*at\s*$', '', title_candidate_from_header, flags=re.IGNORECASE).strip().rstrip(',').strip()
                            
                            if title_candidate_from_header and \
                               (title_candidate_from_header[0].isupper() or title_candidate_from_header[0].isnumeric()) and \
                               title_candidate_from_header.lower() not in COMMON_SECTION_HEADERS and \
                               len(title_candidate_from_header.split()) < 7:
                                new_job_title = title_candidate_from_header
                
                if new_company_name: # Add to global set of companies
                    extracted_companies.add(new_company_name)

                current_role = {
                    "job_title": new_job_title,
                    "company": new_company_name,
                    "start_date": start_date_obj,
                    "end_date": end_date_obj,
                    "description": description_from_date_line # Start with desc from date line
                }
                potential_header_lines = [] # Clear buffer as we've anchored a role

            else: # Line is NOT a date line
                # Could be a standalone header line (Title/Company on its own) or a description line.
                is_this_line_a_new_standalone_header = False
                
                # Heuristic: A new header is usually short, not starting with a bullet, and capitalized.
                # And does not look like a typical description sentence.
                if not sent_text.startswith(('-', '*', '\u2022', '•', '·')) and len(sent_text.split()) < 10:
                    temp_title_sh = None
                    temp_company_sh = None
                    
                    doc_line = kit.nlp(sent_text)
                    # Check for ORG for company
                    for ent in doc_line.ents:
                        if ent.label_ == "ORG" and len(ent.text.split()) < 7:
                            temp_company_sh = ent.text.strip() # Potential company
                            break 
                    
                    # Check for potential title (often the whole line if no company, or part before company)
                    if temp_company_sh:
                        company_idx_sh = sent_text.lower().find(temp_company_sh.lower())
                        if company_idx_sh > 0:
                            cand_sh = sent_text[:company_idx_sh].strip().rstrip(',').strip()
                            if cand_sh and (cand_sh[0].isupper() or cand_sh[0].isnumeric()) and len(cand_sh.split()) < 7:
                                temp_title_sh = cand_sh
                        # If company_idx_sh is 0, the line is just the company name
                    else: # No company found on this line, could it be a title-only line?
                        if (sent_text[0].isupper() or sent_text[0].isnumeric()) and \
                           sent_text.lower() not in COMMON_SECTION_HEADERS and \
                           len(sent_text.split()) < 7:
                            temp_title_sh = sent_text
                    
                    if temp_title_sh or temp_company_sh: # If we found either a title or company
                        is_this_line_a_new_standalone_header = True

                if is_this_line_a_new_standalone_header:
                    # This line looks like a new job's header (Title/Company without a date yet)
                    # Finalize the previous role if it was active and had a start date
                    if current_role.get("start_date"):
                        prev_start_h = current_role.get("start_date")
                        prev_end_h = current_role.get("end_date")
                        if prev_start_h and prev_end_h:
                            try:
                                duration_h = relativedelta(prev_end_h, prev_start_h)
                                total_experience_duration_days += (duration_h.years * 365.25) + (duration_h.months * 30.4) + duration_h.days
                            except TypeError: 
                                logging.warning(f"Could not calculate duration (header): {current_role.get('job_title') or current_role.get('company')}")
                        if current_role.get("job_title") or current_role.get("company") or (prev_start_h and prev_end_h):
                             parsed_resume["experience"].append(current_role)
                             logging.debug(f"===> Appended role (new standalone header found): {current_role.get('job_title') or current_role.get('company')}")
                        current_role = {} # Reset, as this header implies a new role context
                    
                    potential_header_lines.append(sent_text) # Add this line to buffer for next date line
                
                elif current_role.get("start_date"): # Line is not a date, not a new standalone header, APPEND to current role's description
                    current_role["description"] = (current_role.get("description", "") + "\n" + sent_text).strip()
                
                else: # No role active (no start_date yet), and not a date line, and not a clear new header. Add to buffer.
                      # This could be part of a multi-line header before any date is encountered.
                    potential_header_lines.append(sent_text)

        # After the loop, finalize the last current_role if it exists and has data
        if current_role.get("start_date") or current_role.get("job_title") or current_role.get("company"): # Check if there's anything to save
            last_start = current_role.get("start_date")
            last_end = current_role.get("end_date")
            if last_start and last_end: # Only add to duration if dates are valid
                try:
                    duration = relativedelta(last_end, last_start)
                    total_experience_duration_days += (duration.years * 365.25) + (duration.months * 30.4) + duration.days
                except TypeError:
                    logging.warning(f"Could not calculate duration for last role: {current_role.get('job_title') or current_role.get('company')}")
            
            # Append if it has at least a title, company or a valid date pair
            if current_role.get("job_title") or current_role.get("company") or (last_start and last_end):
                 parsed_resume["experience"].append(current_role)
                 logging.debug(f"===> Appended FINAL role: {current_role.get('job_title') or current_role.get('company')}")
            elif potential_header_lines and not parsed_resume["experience"]: # If no roles parsed and buffer has text, maybe it's unstructured
                 logging.warning(f"Experience section had text in potential_header_lines but no structured roles were extracted: {' '.join(potential_header_lines)[:100]}")


        parsed_resume["total_years_experience"] = round(total_experience_duration_days / 365.25, 1)
        logging.info(f"Extracted {len(parsed_resume.get('experience', []))} experience entries. Total calculated years: {parsed_resume.get('total_years_experience',0.0)}")

    else: # No experience_text found
        parsed_resume["total_years_experience"] = 0.0
        if "experience" not in parsed_resume: # Defensive
            parsed_resume["experience"] = []
    
    # This should be outside the "if experience_text:" block, so it's always populated
    parsed_resume["companies"] = sorted(list(extracted_companies))
    
    
    """
    #-------Experience Extraction--------
    total_experience_duration_days = 0
    extracted_companies = set()
    experience_text = sections.get("experience", "")

    if experience_text: 
        logging.info("Parsing experience from 'experience' section.")

        current_role = {}
        potential_title_lines = []

        #we are processing line by line not sentence by sentence bc spacy keeps making mistakes while separating sentences
        for line in experience_text.splitlines():
            sent_text = line.strip()
            if not sent_text: continue

            date_range_pattern =  r"(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|\d{4}|\d{1,2}/\d{4})\s*(?:-|–|to|until)\s*(\b(?:Jan|Feb|Mar|Apr|May|Jun|Jul|Aug|Sep|Oct|Nov|Dec|January|February|March|April|May|June|July|August|September|October|November|December)\s+\d{4}|\d{4}|\d{1,2}/\d{4}|Present|Current|Now)\b"

            date_match = re.search(date_range_pattern,sent_text,re.IGNORECASE)
            
            start_date_obj = None
            end_date_obj = None
            company_name = None
            job_title = None

            if date_match:              
                start_date_str = date_match.group(1)
                end_date_str = date_match.group(2)

                start_date_obj = kit.parse_date(start_date_str)
                end_date_obj = kit.parse_date(end_date_str)

                #if dates are parsed, then we will try to find company title
                if start_date_obj and end_date_obj:
                    for ent in kit.nlp(sent_text).ents:
                        if ent.label_ == "ORG":
                            company_name = ent.text.strip()
                            extracted_companies.add(company_name)
                            break
                    if company_name and job_title is None: # Check if company was found and title is still missing
                        try:
                            # Find where the company name starts in the sentence
                            company_start_index = sent_text.find(company_name)
                            # Check if company name was found and isn't at the very beginning
                            if company_start_index != -1: #it returns -1 if not found
                                text_before_company = sent_text[:company_start_index].strip()
                                text_before_company = re.sub(r'[,|]\s*$', '', text_before_company).strip()
                                if text_before_company and text_before_company[0].isupper():
                                    job_title = text_before_company 
                                    #print(f"--> Guessed title from same line: '{job_title}'") 
                        except Exception as e:
                            logging.warning(f"Error guessing title from same line: {e}")    
                    
                    
                    if potential_title_lines and job_title is None:
                        title_line = potential_title_lines[-1]
                        title_doc = kit.nlp(title_line)
                        for chunk in title_doc.noun_chunks:
                            if chunk.text[0].isupper():
                                job_title = chunk.text
                                break
                        if not job_title and title_doc:
                            #if no chunk worked then we'll use whole line as title guess
                            job_title = title_doc.text

                    if current_role:
                            #if we have this then this means that role ended, now we calculate its duration          
                            prev_start = current_role.get("start_date")
                            prev_end = current_role.get("end_date")
                            if prev_end and prev_start:
                                try:
                                    duration = relativedelta(prev_end,prev_start)
                                    total_experience_duration_days += (duration.years * 365.25) + (duration.months * 30.4) + duration.days
                                except TypeError:
                                    logging.warning(f"Could not calculate duration for role: {current_role.get('job_title')}")
                            
                            logging.debug(f"===> APPENDING final role: {current_role}")
                            parsed_resume["experience"].append(current_role)

                    
                    description_from_date_line = ""

                    try:
                        date_end_index = date_match.end()
                        #Check if there is actually any text *after* the date pattern on this line.
                        # Is the end index of the date match *before* the total length of the line?
                        if date_end_index < len(sent_text):
                            potential_description = sent_text[date_end_index:].strip()
                            # This regex removes one or more hyphens, closing parentheses, asterisks
                            potential_description = re.sub(r'^[-)*\u2022•·\s]+', '', potential_description).strip()
                            if potential_description:
                                description_from_date_line = potential_description
                                logging.debug(f"--> Found description on same line: '{description_from_date_line[:50]}...'")
                    except Exception as e:
                        logging.warning(f"Error extracting description from date line: {e}")


                    current_role = {
                        "job_title":job_title,
                        "company":company_name,
                        "start_date":start_date_obj,
                        "end_date":end_date_obj,
                        "description": description_from_date_line
                    }  
                    potential_title_lines = []
            
            #sentence does not contain a date range
            else:
                is_potential_title = False
                if len(sent_text.split()) < 7 and any(word[0].isupper() for word in sent_text.split() if len(word)>1):
                    is_potential_title = True

                    for ent in kit.nlp(sent_text).ents:
                        if ent.label_ == "ORG":
                            extracted_companies.add(ent.text.strip())
                            is_potential_title = False
                            break


                if is_potential_title:
                    potential_title_lines.append(sent_text)
                
                #print(f"IS POTENTIAL TITLE GUESS: {is_potential_title}")

                if not is_potential_title and current_role:
                    #print(f"APPENDING TO DESCRIPTION for role: {current_role.get('job_title') or current_role.get('company')}") 
                    current_role["description"] = (current_role.get("description", "") + "\n" + sent_text).strip()
                    #print(f"New Description: '{current_role['description'][:50]}...'")

                #elif is_potential_title:
                   # print(f"ADDING TO POTENTIAL TITLE LINES: '{sent_text}'") 
                #elif not current_role:
                    #print("SKIPPING (No current role established yet)")

        if current_role:
            logging.debug(f"\n--- Finalizing Last Role --- LAST ROLE DETAILS: {current_role}")
            last_start = current_role.get("start_date")
            last_end = current_role.get("end_date")

            if last_start and last_end:
                try:
                    duration = relativedelta(last_end, last_start)
                    total_experience_duration_days += (duration.years * 365.25) + (duration.months * 30.4) + duration.days
                except TypeError:
                     logging.warning(f"Could not calculate duration for last role: {current_role.get('job_title')}")
            
            #print(f"===> APPENDING previous role: {current_role}")
            parsed_resume["experience"].append(current_role)

        parsed_resume["total_years_experience"] = round(total_experience_duration_days / 365.25, 1)
        logging.info(f"Extracted {len(parsed_resume['experience'])} experience entries. Total calculated years: {parsed_resume['total_years_experience']}")

        """
    
        

    text_for_contacts = sections.get("header", "")
    if not text_for_contacts: 
         text_for_contacts = "\n".join(sections.values())

    email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
    phone_pattern = r'\b(\+?\d{1,3}[-.\s]?)?\(?\d{3}\)?[-.\s]?\d{3}[-.\s]?\d{4}\b'

    emails = re.findall(email_pattern, text_for_contacts)
    
    if emails:
        parsed_resume["contact_info"]["emails"] = list(set(emails))
        logging.info(f"Found emails: {parsed_resume['contact_info']['emails']}")
    
    phones_found = []
    # Use finditer to get match objects
    for match in re.finditer(phone_pattern, text_for_contacts):
        full_match = match.group(0) 
        cleaned_phone = re.sub(r'[-.\s()]', '', full_match) 
        if cleaned_phone: 
            phones_found.append(cleaned_phone)
    
    if phones_found:
        parsed_resume["contact_info"]["phones"] = list(set(phones_found)) # Keep unique
        logging.info(f"Found phones: {parsed_resume['contact_info']['phones']}")


    # Deduplicate companies extracted from experience
    parsed_resume["companies"] = sorted(list(extracted_companies))

    for section_name, text_content in sections.items():
        if text_content and isinstance(text_content, str): # Ensure content exists and is a string
            doc = kit.nlp(text_content)
            for ent in doc.ents:
                parsed_resume["raw_entities"][ent.label_].append(ent.text.strip())
    
    # Deduplicate all raw entities collected
    for label in parsed_resume["raw_entities"]:
        parsed_resume["raw_entities"][label] = sorted(list(set(parsed_resume["raw_entities"][label])))
    
    return parsed_resume


def parse_resume_file(filepath ,nlp_model_global, tech_skills_list_global, tech_skills_set_global, SECTION_HEADERS_global, EDUCATION_LEVELS_global):
    
    kit = ParsingKit(
        nlp_model=nlp_model_global,
        tech_skills_list_ref=tech_skills_list_global, 
        tech_skills_set_ref=tech_skills_set_global,
        section_headers_ref=SECTION_HEADERS_global,
        education_levels_ref=EDUCATION_LEVELS_global
    )
    
    raw_text = None

    file_extension = os.path.splitext(filepath)[1].lower()

    try:
        file_extension = os.path.splitext(filepath)[1].lower()
        if file_extension == ".txt":
            raw_text = read_text_file(filepath)
        elif file_extension == ".docx":
            raw_text = read_docx_file(filepath)
        elif file_extension == ".pdf":  
            raw_text = read_pdf_file(filepath)
        else:
            logging.error(f"Unsupported file type '{file_extension}'. This function currently only supports .txt files.")
            return {"error": f"Unsupported file type: {file_extension}", "filepath": filepath}

        if raw_text is None or not raw_text.strip():
            logging.error(f"No text extracted or text is empty from file: {filepath}")
            return {"error": "File is empty or could not be read", "filepath": filepath}
        
        cleaned_text = clean_text(raw_text) 

        #debug for doc and pdf

        if file_extension in [".docx", ".pdf"]: 
            print(f"\n--- {file_extension.upper()} Cleaned Text Lines (for Segmentation Debug) ---")
            for i, line in enumerate(cleaned_text.splitlines()):
                if i < 30 : 
                    print(f"Line {i:02d} (len {len(line):03d}): '{line}' --> Repr: {repr(line)}")
                else:
                    break
            print(f"--- End {file_extension.upper()} Cleaned Text Lines ---\n")
        
        #end debug


        sections = segment_resume(cleaned_text, kit.SECTION_HEADERS) 
        
        if not sections:
            logging.error("Segmentation returned no sections.")
            return {"error": "Segmentation failed", "raw_text_snippet": cleaned_text[:200]}

       
        final_dictionary = parse_resume_sections(sections, kit) 

        if final_dictionary and 'experience' in final_dictionary and isinstance(final_dictionary['experience'], list):
            for job in final_dictionary['experience']:
                if isinstance(job, dict):
                    if 'start_date' in job and isinstance(job['start_date'], (datetime.date, datetime.datetime)):
                        job["start_date"] = job['start_date'].strftime('%Y-%m-%d')
                    if 'end_date' in job and hasattr(job['end_date'], 'isoformat'):
                        if isinstance(job['end_date'], (datetime.date, datetime.datetime)):
                            job['end_date'] = job['end_date'].strftime('%Y-%m-%d')
        
        if cleaned_text and final_dictionary:
             final_dictionary["raw_text_snippet"] = cleaned_text[:750] 

        return final_dictionary

    except FileNotFoundError:
        logging.error(f"Error: File not found at {filepath}")
        return {"error": f"File not found during parse_resume_file: {filepath}", "filepath": filepath}
    except Exception as e:
        logging.error(f"General error parsing resume file {filepath}: {e}")
        return {"error": f"General error during parse_resume_file: {str(e)}", "filepath": filepath}
    

def get_text_from_txt_object(uploaded_file_object):
    try:
        raw_text = uploaded_file_object.read().decode('utf-8')
        logging.info("Successfully read TXT file content.")
        return raw_text
    except Exception as e:
        logging.error(f"Error reading TXT file object: {e}")
        return None

def get_text_from_docx_object(uploaded_file_object):
    try:
        doc = Document(uploaded_file_object) 
        full_text = [para.text for para in doc.paragraphs]
        raw_text = '\n'.join(full_text)
        logging.info("Successfully read DOCX file content.")
        return raw_text
    except Exception as e:
        logging.error(f"Error reading DOCX file object: {e}")
        return None

def get_text_from_pdf_object(uploaded_file_object):
    try:
        full_text = []

        if hasattr(uploaded_file_object, 'seek') and callable(uploaded_file_object.seek):
            uploaded_file_object.seek(0)

        with pdfplumber.open(uploaded_file_object) as pdf:
            for page in pdf.pages:
                extracted_page_text = page.extract_text()
                if extracted_page_text: 
                    full_text.append(extracted_page_text)
        raw_text = '\n'.join(full_text)
        if not raw_text.strip():
            file_name_attr = getattr(uploaded_file_object, 'name', 'Uploaded PDF Object')
            logging.warning(f"No text extracted from PDF: {file_name_attr}")
            return None 
        logging.info("Successfully read PDF file content.")
        return raw_text
    except Exception as e:
        file_name_attr = getattr(uploaded_file_object, 'name', 'Uploaded PDF Object')
        logging.error(f"Error reading PDF file object {file_name_attr}: {e}")
        return None


#Deals with Streamlit's UploadedFile objects (which could be TXT, DOCX, PDF from memory).
def process_streamlit_file(
        uploaded_file_object,
        nlp_ref,
        tech_skills_list_ref, 
        tech_skills_set_ref, 
        section_headers_ref, 
        education_levels_ref
):
    logging.info(f"--- process_streamlit_file: STARTED for {uploaded_file_object.name} ---")
    file_type = uploaded_file_object.type
    raw_text = None

    if file_type == "text/plain":
        raw_text = get_text_from_txt_object(uploaded_file_object)
    elif file_type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document": # DOCX
        raw_text = get_text_from_docx_object(uploaded_file_object)
    elif file_type == "application/pdf":
        raw_text = get_text_from_pdf_object(uploaded_file_object)
    else:
        logging.error(f"Unsupported file type: {file_type}")
        return None

    if raw_text is None or not raw_text.strip():
        logging.error("No text extracted or text is empty from streamlit file.")
        return None

    logging.info(f"Raw text extracted for parsing (snippet): {raw_text[:750]}...")


    kit = ParsingKit(
        nlp_model=nlp_ref,
        tech_skills_list_ref=tech_skills_list_ref,
        tech_skills_set_ref=tech_skills_set_ref,
        section_headers_ref=section_headers_ref,
        education_levels_ref=education_levels_ref
    )
    logging.info("ParsingKit initialized.")

    cleaned_text = clean_text(raw_text)
    sections = segment_resume(cleaned_text, kit.SECTION_HEADERS) 
    logging.info(f"Text segmented into sections: {list(sections.keys())}")


    if not sections:
        logging.error("Segmentation returned no sections.")
        return {"error": "Segmentation failed", "raw_text_snippet": raw_text[:200]}

    final_dictionary = parse_resume_sections(sections,kit) 
    logging.info("Detailed parsing of sections complete.")

    if final_dictionary is None:
        logging.error("parse_resume_sections returned none, which is unexpected.")
        return {"error": "Parsing failed unexpectedly"}

    # Converting datetime objects to strings
    if 'experience' in final_dictionary and isinstance(final_dictionary['experience'], list):
        for job in final_dictionary['experience']:
            if isinstance(job, dict):
                if 'start_date' in job and isinstance(job['start_date'], (datetime.date, datetime.datetime)):
                    job["start_date"] = job['start_date'].strftime('%Y-%m-%d')
                if 'end_date' in job and hasattr(job['end_date'], 'isoformat'):
                    if isinstance(job['end_date'], (datetime.date, datetime.datetime)):
                        job['end_date'] = job['end_date'].strftime('%Y-%m-%d')

    
    if cleaned_text and final_dictionary: 
        if "raw_text_snippet" not in final_dictionary:
            final_dictionary["raw_text_snippet"] = cleaned_text[:200]

    logging.info(f"--- process_streamlit_file: RETURNING (snippet): {str(final_dictionary)[:300]} ---")
    return final_dictionary



#-------------------------------------------------------------------------------


"""
if __name__ == '__main__':

    SECTION_HEADERS_CONFIG_MAIN = {
        "summary": r"^\s*(summary|profile|objective|about\s*me)\s*[:\n]",
        "skills": r"^\s*(skills|skills\s*&\s*abilities|abilities|technical\s*skills|technical\s*proficiency|core\s*competencies|technologies)([:\s]|\s*$)",
        "experience": r"^\s*(experience|work\s*experience|employment\s*history|professional\s*experience)\s*[:\n]",
        "education": r"^\s*(education|academic\s*background|academic\s*qualifications)\s*[:\n]",
        "projects": r"^\s*(projects|personal\s*projects)\s*[:\n]",
    }

    EDUCATION_LEVELS_CONFIG_MAIN = { # Renamed to avoid conflict
        "phd": 5, "doctorate": 5, "dphil": 5, "doctor of philosophy": 5,
        "master": 4, "masters": 4, "msc": 4, "meng": 4, "mba": 4, "ma": 4, "ms": 4, "master of science": 4, "master of arts": 4,
        "bachelor": 3, "bachelors": 3, "bsc": 3, "beng": 3, "ba": 3, "bs": 3, "b.s":3, "b.a":3, "b.sc":3,
        "associate": 2, "associates": 2, "associate degree": 2,
        "college": 1, "diploma": 1, "advanced diploma":1,
        "high school": 0, "ged": 0, "secondary school": 0
    }
    
     
    nlp_model_main = NLP_MODEL_GLOBAL 
    if nlp_model_main is None:
        logging.critical("NLP Model (NLP_MODEL_GLOBAL) is None. Cannot proceed with parsing in main. Exiting.")
        exit()
 
    skills_list = TECH_SKILLS_LIST_GLOBAL
    skills_set = TECH_SKILLS_SET_GLOBAL
    SECTION_HEADERS_CONFIG = SECTION_HEADERS_GLOBAL 
    EDUCATION_LEVELS_CONFIG = EDUCATION_LEVELS_GLOBAL


    script_dir = os.path.dirname(os.path.abspath(__file__))
    dummy_resume_filepath = os.path.join(script_dir, 'tests', 'data', 'raw_resumes', "resume_01.txt") 
    dummdummy_docx_filepath = os.path.join(script_dir, 'tests', 'data', 'raw_resumes', "resume_11.docx") 
    skills_filepath = os.path.join(script_dir, 'tests', 'data','skills.json') 
    dummy_pdf_filepath = os.path.join(script_dir, 'tests', 'data', 'raw_resumes', "resume_17.pdf")

    tech_skills_list_main = skills_list
    tech_skills_set_main = skills_set

    test_data_output_dir = os.path.join(script_dir, 'test_data_output_files')
    os.makedirs(test_data_output_dir, exist_ok=True)

    logging.info(f"--- Attempting to parse: {dummy_pdf_filepath} ---")
   

    if os.path.exists(dummy_pdf_filepath):
        logging.info(f"Attempting to read DOCX file directly: {dummy_pdf_filepath}")
        parsed_info_docx = parse_resume_file(
        dummy_pdf_filepath, 
        nlp_model_main,                   # Your globally loaded nlp model
        tech_skills_list_main,           # Your globally loaded LIST of skills
        tech_skills_set_main,       # Your globally loaded SET of skills
        SECTION_HEADERS_CONFIG,       # Your globally defined SECTION_HEADERS
        EDUCATION_LEVELS_CONFIG       # Your globally defined EDUCATION_LEVELS
    )
        if parsed_info_docx and "error" not in parsed_info_docx:
            output_json_path_docx = os.path.join(test_data_output_dir, "parsed_resume_docx.json")
            with open(output_json_path_docx, 'w', encoding='utf-8') as outfile_docx:
                json.dump(parsed_info_docx, outfile_docx, indent=4)
            logging.info(f"DOCX Parsed data saved to: {output_json_path_docx}")
            print("\n--- Parsed DOCX Data ---")
            print(json.dumps(parsed_info_docx, indent=4))
        elif parsed_info_docx:
            logging.error(f"Error parsing DOCX resume { dummy_pdf_filepath}: {parsed_info_docx.get('error')}")
        else:
            logging.error(f"Failed to parse DOCX resume (returned None): { dummy_pdf_filepath}")
    else:
        logging.warning(f"Test DOCX file not found at {dummy_pdf_filepath}, skipping DOCX parsing test.")
        

"""